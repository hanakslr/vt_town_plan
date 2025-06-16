"""
Given a docx file, extract its content to a structured JSON file to be used for chunking
and generating embeddings.
"""

import json
import os
from collections import defaultdict
from dataclasses import asdict, dataclass, is_dataclass
from typing import Any, Dict, List, Optional
from zipfile import ZipFile

import psycopg
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

from elements import (
    Action,
    ActionTable,
    Caption,
    DocumentSection,
    Fact,
    Goals2050,
    Heading,
    Objective,
    PublicEngagementFindings,
    Strategy,
    StructuredDocument,
    ThreeFacts,
)
from elements import (
    Paragraph as Para,
)
from elements import (
    Table as TableModel,
)
from parsers.table_parsers import TableMerger, TableParserFactory, TableStyles


# Step 1: Load the numbering.xml file from the .docx archive
def load_numbering_xml(docx_path):
    with ZipFile(docx_path) as docx_zip:
        numbering_xml = docx_zip.read("word/numbering.xml")
        return etree.fromstring(numbering_xml)


# Step 2: Parse numbering definitions from numbering.xml
def parse_numbering_definitions(numbering_tree):
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    numId_to_abstractId = {}
    abstractId_to_levels = {}

    for num in numbering_tree.findall(".//w:num", namespaces=ns):
        numId = num.get(f"{{{ns['w']}}}numId")
        abstractNumId_el = num.find(".//w:abstractNumId", namespaces=ns)
        if abstractNumId_el is not None:
            numId_to_abstractId[numId] = abstractNumId_el.get(f"{{{ns['w']}}}val")

    for abstractNum in numbering_tree.findall(".//w:abstractNum", namespaces=ns):
        abstractId = abstractNum.get(f"{{{ns['w']}}}abstractNumId")
        levels = {}
        for lvl in abstractNum.findall(".//w:lvl", namespaces=ns):
            ilvl = lvl.get(f"{{{ns['w']}}}ilvl")
            numFmt = lvl.find("w:numFmt", namespaces=ns)
            lvlText = lvl.find("w:lvlText", namespaces=ns)
            levels[ilvl] = {
                "numFmt": numFmt.get(f"{{{ns['w']}}}val")
                if numFmt is not None
                else None,
                "lvlText": lvlText.get(f"{{{ns['w']}}}val")
                if lvlText is not None
                else None,
            }
        abstractId_to_levels[abstractId] = levels

    return numId_to_abstractId, abstractId_to_levels


@dataclass
class Image:
    """Represents an image extracted from the document."""

    filename: str
    alt_text: Optional[str] = None
    section_path: Optional[List[str]] = None
    width: Optional[int] = None
    height: Optional[int] = None


class DocumentExtract:
    doc: Document
    current_section: list[Heading]
    numbering_tree: Any
    document: StructuredDocument
    images: List[Image]
    image_output_dir: str

    def __init__(self, file_path: str, image_output_dir: str = "output/images") -> None:
        self.doc = Document(file_path)
        self.current_section = []
        self.document = StructuredDocument()
        self.numbering_tree = load_numbering_xml(file_path)
        self.numId_to_abstractId, self.abstractId_to_levels = (
            parse_numbering_definitions(self.numbering_tree)
        )
        self.list_counters = defaultdict(lambda: defaultdict(int))
        self.images = []
        self.image_output_dir = image_output_dir
        os.makedirs(image_output_dir, exist_ok=True)

    def extract(self):
        """
        Process the document, extracting the elements into a structured document.
        """
        general_content = []
        previous_table = None

        for block in self._iter_block_items():
            if isinstance(block, Paragraph):
                data = self._extract_paragraph(block)
                if data and getattr(data, "text", None):  # Skip empty paragraphs
                    self._add_to_document(data, general_content)
            elif isinstance(block, Table):
                data = self._extract_table(block)
                if data:
                    # Check if we should merge with a previous table
                    if (
                        previous_table
                        and hasattr(previous_table, "type")
                        and hasattr(data, "type")
                        and TableMerger.should_merge(
                            self._serialize_element(previous_table),
                            self._serialize_element(data),
                        )
                    ):
                        # Merge with previous table instead of adding a new one
                        previous_table = TableMerger.merge_tables(
                            self._serialize_element(previous_table),
                            self._serialize_element(data),
                        )
                        # Convert merged dict back to dataclass
                        previous_table = self._dict_to_element(previous_table)

                        # Update the appropriate section of the document
                        self._update_merged_table(previous_table, general_content)
                    else:
                        self._add_to_document(data, general_content)
                        previous_table = data
            else:
                raise Exception(f"Unexpected instance item {block}")

        # Update the document's content with the processed general content
        self.document.content = general_content

    def _add_to_document(self, data, general_content):
        """Add an extracted element to the appropriate place in the structured document."""
        # Main chapter heading
        if isinstance(data, Heading) and data.level == 1:
            self.document.chapter_number = getattr(data, "chapter_number", None)
            self.document.title = data.text

        # 2050 goals
        elif isinstance(data, Goals2050):
            self.document.goals_2050 = data

        # Three facts
        elif isinstance(data, ThreeFacts):
            self.document.three_facts = data

        # Public engagement findings
        elif isinstance(data, PublicEngagementFindings):
            self.document.public_engagement = data

        # Action table
        elif isinstance(data, ActionTable):
            self.document.actions = data

        # All other content - paragraphs, sections, captions, etc.
        else:
            # Keep the section_path property for organizing content
            # It helps with preserving the document structure
            general_content.append(data)

    def _update_merged_table(self, merged_table, general_content):
        """Update the document with a merged table."""
        # Action table
        if isinstance(merged_table, ActionTable):
            self.document.actions = merged_table
        # 2050 goals
        elif isinstance(merged_table, Goals2050):
            self.document.goals_2050 = merged_table
        # Three facts
        elif isinstance(merged_table, ThreeFacts):
            self.document.three_facts = merged_table
        # Public engagement findings
        elif isinstance(merged_table, PublicEngagementFindings):
            self.document.public_engagement = merged_table
        # Other table types
        else:
            # Find and replace in general content
            for i in range(len(general_content) - 1, -1, -1):
                if isinstance(general_content[i], type(merged_table)):
                    general_content[i] = merged_table
                    break

    def encode(self):
        # Will be updated as per your requirements
        pass

    def dump_to_file(self, path):
        """Save the structured document to a JSON file."""
        # Add images to the document structure
        self.document.images = [asdict(img) for img in self.images]

        with open(path, "w", encoding="utf-8") as f:
            json.dump(self.document.to_dict(), f, indent=2)

    def dump_to_db(self):
        """Save the structured document to the database."""
        conn = psycopg.connect(
            dbname="db", user="admin", password="password", host="localhost", port=5431
        )
        cur = conn.cursor()

        doc_dict = self.document.to_dict()

        # Insert chapter metadata
        cur.execute(
            """
            INSERT INTO chapters (chapter_number, title)
            VALUES (%s, %s)
            RETURNING id
            """,
            (
                doc_dict.get("chapter_number"),
                doc_dict.get("title"),
            ),
        )
        chapter_id = cur.fetchone()[0]

        # Insert special sections if they exist
        if "2050_goals" in doc_dict:
            cur.execute(
                """
                INSERT INTO goals (chapter_id, content)
                VALUES (%s, %s)
                """,
                (chapter_id, json.dumps(doc_dict["2050_goals"])),
            )

        if "actions" in doc_dict:
            cur.execute(
                """
                INSERT INTO actions (chapter_id, content)
                VALUES (%s, %s)
                """,
                (chapter_id, json.dumps(doc_dict["actions"])),
            )

        # Insert content chunks
        for item in doc_dict.get("content", []):
            item_type = item.get("type", "unknown")
            item_text = item.get("text", "")

            cur.execute(
                """
                INSERT INTO plan_chunks (chapter_id, content_type, content)
                VALUES (%s, %s, %s)
                """,
                (
                    chapter_id,
                    item_type,
                    item_text,
                ),
            )

        conn.commit()
        cur.close()
        conn.close()

    def _iter_block_items(self):
        """Yield paragraphs and tables in document order"""
        parent_elm = self.doc._element.body
        for child in parent_elm.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, self.doc)
            elif child.tag == qn("w:tbl"):
                yield Table(child, self.doc)
            elif child.tag == qn("w:sectPr"):
                pass
            else:
                raise Exception(f"Unhandled tag -  {child.tag}")

    # Track list state and render list numbers
    def build_list_number(self, numId, ilvl):
        abstract_id = self.numId_to_abstractId[numId]
        levels = self.abstractId_to_levels[abstract_id]

        lvlText = levels[str(ilvl)]["lvlText"]
        label = lvlText

        ilvl = int(ilvl)
        for lvl in range(ilvl + 1, 9):
            self.list_counters[numId][lvl] = 0
        self.list_counters[numId][ilvl] += 1

        for n in range(1, 10):
            if f"%{n}" in label:
                lvl_index = n - 1
                counter = self.list_counters[numId].get(lvl_index, 0)
                label = label.replace(f"%{n}", str(counter))

        return label

    def _extract_table(self, table: Table) -> DocumentSection:
        rows = []
        styles = []
        seen_cells = set()
        for row in table.rows:
            cols = []
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                text = cell.text.strip()
                style_info = None
                for para in cell.paragraphs:
                    for run in para.runs:
                        font = run.font
                        style_info = TableStyles(
                            text=run.text,
                            font_name=font.name,
                            font_size=font.size.pt if font.size else None,
                            bold=font.bold,
                            italic=font.italic,
                        )
                        break
                    if style_info:
                        break
                if text:
                    cols.append(text)
                    if style_info:
                        styles.append(style_info)
            if cols:
                rows.append(cols)
        if not rows:
            return None
        section_texts = [s.text for s in self.current_section]
        parser = TableParserFactory.create_parser(
            table=table,
            rows=rows,
            styles=styles,
            current_section=section_texts,
            list_number_generator=self.build_list_number,
        )
        result = parser.parse()
        if (
            result
            and hasattr(result, "type")
            and result.type == "heading"
            and hasattr(result, "level")
            and result.level == 1
        ):
            self.current_section.append(
                Heading(
                    level=1,
                    text=result.text,
                    chapter_number=getattr(result, "chapter_number", None),
                )
            )
        return result if result else None

    def _extract_images_from_paragraph(self, paragraph: Paragraph) -> List[Image]:
        """Extract images from a paragraph."""
        images = []
        for run in paragraph.runs:
            for element in run._element.iterchildren():
                if element.tag == qn("w:drawing"):
                    # Find the blip element which contains the image reference
                    blip = element.find(
                        ".//a:blip",
                        {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"},
                    )
                    if blip is not None:
                        # Get the image ID from the blip element
                        embed = blip.get(qn("r:embed"))
                        if embed:
                            # Get the image data from the document
                            image_part = self.doc.part.related_parts[embed]
                            # Generate a unique filename
                            image_filename = f"image_{len(self.images) + 1}.{image_part.content_type.split('/')[-1]}"
                            image_path = os.path.join(
                                self.image_output_dir, image_filename
                            )

                            # Save the image
                            with open(image_path, "wb") as f:
                                f.write(image_part.blob)

                            # Get image dimensions if available
                            width = None
                            height = None
                            extent = element.find(
                                ".//wp:extent",
                                {
                                    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                                },
                            )
                            if extent is not None:
                                width = (
                                    int(extent.get("cx", 0)) // 9525
                                )  # Convert EMUs to pixels
                                height = int(extent.get("cy", 0)) // 9525

                            # Create image object
                            image = Image(
                                filename=image_filename,
                                alt_text=paragraph.text.strip()
                                if paragraph.text.strip()
                                else None,
                                section_path=[s.text for s in self.current_section],
                                width=width,
                                height=height,
                            )
                            images.append(image)
                            self.images.append(image)
        return images

    def _extract_paragraph(self, paragraph: Paragraph) -> DocumentSection:
        if not paragraph.text.strip() and not any(
            run._element.findall(qn("w:drawing")) for run in paragraph.runs
        ):
            return None

        # Extract any images in the paragraph
        images = self._extract_images_from_paragraph(paragraph)

        if paragraph.style.name == "Section Heading":
            level = 2
            while self.current_section and self.current_section[-1].level >= level:
                self.current_section.pop()
            curr_heading = Heading(level=level, text=paragraph.text.strip())
            self.current_section.append(curr_heading)
            return Heading(
                level=curr_heading.level,
                text=curr_heading.text,
                section=self.current_section[0].text,
            )
        elif paragraph.style.name == "List Paragraph":
            return Para(
                paragraph_style=paragraph.style.name,
                text=paragraph.text.strip(),
                section_path=[s.text for s in self.current_section],
                images=images if images else None,
            )
        elif paragraph.style.name in ["Normal", "No Spacing", "paragraph"]:
            return Para(
                paragraph_style=paragraph.style.name,
                text=paragraph.text.strip(),
                section_path=[s.text for s in self.current_section],
                images=images if images else None,
            )
        elif paragraph.style.name == "Caption":
            return Caption(
                text=paragraph.text.strip(),
                section_path=[s.text for s in self.current_section],
                images=images if images else None,
            )
        else:
            raise Exception(
                f"Unknown paragraph style {paragraph.style.name}. Content: {paragraph.text.strip()}"
            )

    def _dict_to_element(self, d: dict) -> DocumentSection:
        if not d or "type" not in d:
            return d
        t = d["type"]
        if t == "heading":
            return Heading(**d)
        elif t == "2050_goals":
            return Goals2050(**d)
        elif t == "3_facts":
            facts = [Fact(**f) for f in d["facts"]]
            return ThreeFacts(text=d["text"], section=d["section"], facts=facts)
        elif t == "3_public_engagement_findings":
            facts = [Fact(**f) for f in d["facts"]]
            return PublicEngagementFindings(
                text=d["text"], section=d["section"], facts=facts
            )
        elif t == "action_table":
            objectives = [Objective(**o) for o in d["objectives"]]
            strategies = [
                Strategy(
                    label=s["label"],
                    text=s["text"],
                    actions=[Action(**a) for a in s["actions"]],
                )
                for s in d["strategies"]
            ]
            return ActionTable(
                section=d["section"], objectives=objectives, strategies=strategies
            )
        elif t == "table":
            return TableModel(
                rows=d["rows"],
                border_info=d.get("border_info"),
                border_color=d.get("border_color"),
            )
        else:
            return d

    def _serialize_element(self, element: Any) -> Dict[str, Any]:
        """Serialize an element to a dictionary using the appropriate method."""
        if isinstance(element, DocumentSection):
            return element.to_dict()
        elif is_dataclass(element):
            return asdict(element)
        elif isinstance(element, dict):
            return element
        else:
            return {"value": str(element)}


# Example usage: `python extract.py files/docx_test.docx`
if __name__ == "__main__":
    import argparse
    from pathlib import Path

    parser = argparse.ArgumentParser(
        description="Extract content from a Word document."
    )
    parser.add_argument(
        "input_file", help="Path to the input Word document (.docx file)"
    )
    args = parser.parse_args()

    # Create output directory if it doesn't exist
    output_dir = Path("output")
    output_dir.mkdir(exist_ok=True)

    # Generate output filename based on input filename
    input_filename = Path(args.input_file).stem
    output_file = output_dir / f"{input_filename}.json"

    ex = DocumentExtract(args.input_file)

    ex.extract()
    # ex.encode()
    ex.dump_to_file(output_file)
    # ex.dump_to_db()

    print(f"Content extracted and saved to: {output_file}")
